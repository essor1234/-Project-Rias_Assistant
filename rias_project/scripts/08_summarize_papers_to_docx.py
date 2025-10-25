#!/usr/bin/env python
# -*- coding: utf-8 -*-

import json
import sys
import datetime
import time
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------
# Logging setup helper
# ---------------------------------------------------------------------
class Tee:
    def __init__(self, *files): self.files = files
    def write(self, obj):
        for f in self.files:
            f.write(obj); f.flush()
    def flush(self):
        for f in self.files: f.flush()


# ---------------------------------------------------------------------
# Main class
# ---------------------------------------------------------------------
class PaperSummarizer:
    def __init__(
        self,
        txt_dir,
        images_dir,
        prompt_path,
        output_path,
        model="gpt-5",
        max_tokens=20000,
        max_retries=3,
    ):
        load_dotenv()
        self.client = OpenAI()
        self.model = model
        self.max_tokens = max_tokens
        self.max_retries = max_retries

        self.txt_dir = Path(txt_dir)
        self.images_dir = Path(images_dir)
        self.prompt_path = Path(prompt_path)
        self.output_path = Path(output_path)
        self.output_path.parent.mkdir(parents=True, exist_ok=True)

        # Setup logging
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        script_dir = Path(__file__).resolve().parent.parent
        log_dir = script_dir / "logs"
        log_dir.mkdir(exist_ok=True)
        log_file = log_dir / f"summary_log_{timestamp}.txt"
        log_f = open(log_file, "w", encoding="utf-8")
        sys.stdout = Tee(sys.__stdout__, log_f)
        sys.stderr = Tee(sys.__stderr__, log_f)
        self.log_file = log_file

    # ------------------------------------------------------------------
    # Utility methods
    # ------------------------------------------------------------------
    def truncate_text(self, text: str, limit: int = 30_000) -> str:
        return text if len(text) <= limit else text[:limit] + "\n\n[Text truncated for LLM]"

    def load_prompt(self) -> str:
        return self.prompt_path.read_text(encoding="utf-8")

    def clean_raw(self, raw: str) -> str:
        """Clean malformed JSON response from LLM."""
        try:
            # Basic string cleaning
            raw = raw.strip()
            
            # If it's already valid JSON, return it
            try:
                json.loads(raw)
                return raw
            except:
                pass

            # Remove any markdown code block markers
            raw = raw.replace("```json", "").replace("```", "")
            
            # Ensure it starts/ends with curly braces
            if not raw.startswith("{"): raw = "{" + raw
            if not raw.endswith("}"): raw = raw + "}"
            
            # Fix common JSON formatting issues
            raw = raw.replace('""', '"')
            raw = raw.replace('}"', '}')
            raw = raw.replace('"{', '{')
            raw = raw.replace('\n', ' ')
            
            # Add quotes around property names if missing
            import re
            raw = re.sub(r'([{,])\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*:', r'\1"\2":', raw)
            
            # If no SummaryDoc key exists, wrap the entire content
            if '"SummaryDoc"' not in raw:
                raw = raw.replace("{", '{"SummaryDoc":', 1)
                
            return raw
            
        except Exception as e:
            print(f"Failed to clean JSON: {e}")
            # Return a minimal valid JSON as fallback
            return '{"SummaryDoc": "Error parsing LLM response"}'

    # ------------------------------------------------------------------
    # LLM call
    # ------------------------------------------------------------------
    def call_llm(self, prompt: str) -> str:
        """Call GPT model and ensure valid JSON output."""
        system_msg = {
            "role": "system",
            "content": (
                "You are an academic summarizer and document restorer. "
                "You MUST output only valid JSON with this schema:\n"
                "{ \"SummaryDoc\": \"<full reconstructed academic document text including figure markers>\" }"
            ),
        }

        user_msg = {"role": "user", "content": prompt}

        for attempt in range(self.max_retries):
            try:
                resp = self.client.chat.completions.create(
                    model=self.model,
                    messages=[system_msg, user_msg],
                    max_completion_tokens=self.max_tokens,
                    response_format={"type": "json_object"},
                )

                choice = resp.choices[0]
                content = getattr(choice.message, "content", None)
                if not content:
                    print("⚠️ Empty content field from LLM.")
                    continue

                cleaned = self.clean_raw(content)
                return cleaned.strip()

            except Exception as e:
                print(f"Attempt {attempt+1}/{self.max_retries} failed: {e}")
                if attempt == self.max_retries - 1:
                    raise
                time.sleep(2 ** attempt)
        return ""

    # ------------------------------------------------------------------
    # DOCX helpers
    # ------------------------------------------------------------------
    def ensure_caption_style(self, doc: Document):
        if "Caption" not in doc.styles:
            style = doc.styles.add_style("Caption", 1)
            style.font.name = "Times New Roman"
            style.font.size = Pt(10)
            style.font.italic = True

    def insert_image(self, doc: Document, img_path: Path, caption: str):
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img.add_run()
            run.add_picture(str(img_path), width=Inches(5.5))

            p_cap = doc.add_paragraph(caption, style="Caption")
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run_cap in p_cap.runs:
                run_cap.italic = True
            doc.add_paragraph()
        except Exception as e:
            print(f"Could not insert image {img_path.name}: {e}")

    def create_docx(self, summary_text: str):
        doc = Document()
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)

        self.ensure_caption_style(doc)

        image_map = {
            p.name.lower(): p for p in self.images_dir.iterdir()
            if p.suffix.lower() in {".png", ".jpg", ".jpeg", ".tif", ".bmp", ".svg"}
        }

        print(f"Found {len(image_map)} images in {self.images_dir}")

        blocks = [b.strip() for b in summary_text.split("\n\n") if b.strip()]

        for block in blocks:
            line = block.strip()

            if line.startswith("###"):
                doc.add_heading(line.lstrip("# ").strip(), level=3)
            elif line.startswith("##"):
                doc.add_heading(line.lstrip("# ").strip(), level=2)
            elif line.startswith("#"):
                doc.add_heading(line.lstrip("# ").strip(), level=1)
            elif line.startswith("[[FIGURE:"):
                try:
                    inner = line.strip("[]").replace("FIGURE:", "").strip()
                    parts = [p.strip() for p in inner.split("|")]
                    filename = parts[0]
                    caption = " | ".join(parts[1:])
                    img_key = Path(filename).name.lower()

                    if img_key in image_map:
                        self.insert_image(doc, image_map[img_key], caption)
                        print(f"Inserted image: {filename}")
                    else:
                        print(f"Image NOT FOUND: {filename}")
                        p = doc.add_paragraph(f"[Image missing: {filename}] {caption}", style="Normal")
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Figure parsing error: {e}\n   Block: {line}")
                    doc.add_paragraph(line, style="Normal")
            else:
                doc.add_paragraph(line, style="Normal")

        doc.save(self.output_path)
        print(f"💾 DOCX saved → {self.output_path}")

    # ------------------------------------------------------------------
    # Main runner
    # ------------------------------------------------------------------
    def run(self, pdf_stem: str):
        """Main process to generate summary document."""
        try:
            print(f"\nProcessing {pdf_stem}")
            
            # Load and process text
            txt_files = sorted(self.txt_dir.glob("*.txt"))
            if not txt_files:
                raise FileNotFoundError(f"No .txt files in {self.txt_dir}")
                
            combined = "\n\n".join(
                f"--- FILE: {f.name} ---\n{f.read_text(encoding='utf-8')}"
                for f in txt_files
            )
            
            # Call LLM with retry on JSON error
            prompt = self.load_prompt().replace("<<<DOCUMENT_TEXT>>>", combined)
            raw = self.call_llm(prompt)
            
            if not raw:
                raise ValueError("Empty response from LLM")
                
            # Parse response with fallback cleaning
            try:
                parsed = json.loads(raw)
            except json.JSONDecodeError:
                cleaned = self.clean_raw(raw)
                try:
                    parsed = json.loads(cleaned)
                except json.JSONDecodeError as e:
                    print(f"Failed to parse JSON even after cleaning: {e}")
                    # Create minimal valid response
                    parsed = {"SummaryDoc": raw}
            
            summary = parsed.get("SummaryDoc", "").strip()
            if not summary:
                raise ValueError("Empty summary document")
                
            # Create final DOCX
            self.create_docx(summary)
            return True
            
        except Exception as e:
            print(f"ERROR in summarize: {e}")
            return False

# ...existing code...

# This function should REPLACE the old `def run(...)` 
# at the end of your 'scripts/01_extract_text.py' file.

# ------------------------------------------------------------------
#     # Main runner
#     # ------------------------------------------------------------------
#     # ... (End of the PaperSummarizer class) ...


# ---------------------------------------------------------------------
# Bridge function for main.py
# ---------------------------------------------------------------------

def run(pdf_path, out_dir, prev=None):
    """
    Bridge used by main.py to run the PaperSummarizer.
    """
    try:
        p = Path(pdf_path)
        out = Path(out_dir)
        pdf_stem = p.stem
        
        # Get base paths
        result_dir = out.parent.parent  # Go up to the PDF's result folder
        processed_dir = result_dir / "processed"
        
        # Input directories from previous pipeline steps
        txt_dir = processed_dir / "01_extract_text_output"
        images_dir = processed_dir / "06_extract_images_output"
        
        # Get project paths
        SCRIPT_DIR = Path(__file__).resolve().parent.parent
        prompt_path = SCRIPT_DIR / "prompts" / "[Prompt]summarize_papers.txt"
        
        # Output file path
        output_docx = out / f"{pdf_stem}_Summary.docx"
        
        # Debug print paths
        print(f"\nInput paths for {pdf_stem}:")
        print(f"- Text dir: {txt_dir}")
        print(f"- Images dir: {images_dir}")
        print(f"- Output DOCX: {output_docx}")
        
        # Verify inputs exist
        if not txt_dir.exists() or not any(txt_dir.glob("*.txt")):
            raise FileNotFoundError(f"No text files found in: {txt_dir}")
            
        if not images_dir.exists():
            print(f"Warning: Images directory not found: {images_dir}")
        
        # Initialize and run summarizer
        summarizer = PaperSummarizer(
            txt_dir=txt_dir,
            images_dir=images_dir,
            prompt_path=prompt_path,
            output_path=output_docx,
            model="gpt-4-turbo",
            max_tokens=4096,
            max_retries=3
        )
        
        if summarizer.run(pdf_stem):
            return {
                "status": "success",
                "files": [output_docx.name],
                "summary": "Summary DOCX created."
            }
        else:
            return {
                "status": "error",
                "error": "Failed to generate summary"
            }
            
    except Exception as e:
        print(f"ERROR in 08_summarize: {e}")
        return {"status": "error", "error": str(e)}

# ----------------------------------------------------------------------
# Optional: CLI entry point (if you want to run this file directly)
# ----------------------------------------------------------------------
if __name__ == "__main__":
    # This is just for testing this script directly
    # You would need to manually set up the paths
    
    print("Running PaperSummarizer in standalone mode...")
    
    # --- Example paths for direct testing ---
    # You MUST change these to match your test setup
    
    PDF_STEM = "test4" # The name of the PDF you want to process
    
    # Assumes your script is in 'rias_project/scripts/'
    ROOT = Path(__file__).resolve().parent.parent 
    
    # MOCK PATHS from a pipeline run
    MOCK_TXT_DIR = ROOT / "results/TEST_SESSION/test4/processed/01_extract_text_output"
    MOCK_IMG_DIR = ROOT / "results/TEST_SESSION/test4/processed/06_extract_images_output"
    MOCK_OUT_FILE = ROOT / "data/summarize_to_doc_output/Direct_Test_Summary.docx"
    
    MOCK_PROMPT = ROOT / "prompts" / "[Prompt]summarize_papers.txt"

    # --- End Example Paths ---
    
    try:
        summarizer = PaperSummarizer(
            txt_dir=MOCK_TXT_DIR,
            images_dir=MOCK_IMG_DIR,
            prompt_path=MOCK_PROMPT,
            output_path=MOCK_OUT_FILE,
            model="gpt-4-turbo",
            max_tokens=4096
        )
        
        summarizer.run(pdf_stem=PDF_STEM)
        
    except Exception as e:
        print(f"Error during standalone test: {e}")
        import traceback
        traceback.print_exc()
# ----------------------------------------------------------------------
# ---------------------------------------------------------------------
#=============================================================
# #!/usr/bin/env python
# # -*- coding: utf-8 -*-

# import json
# import sys
# import datetime
# import time
# from pathlib import Path
# from dotenv import load_dotenv
# from openai import OpenAI
# from docx import Document
# from docx.shared import Pt, Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH

# # ---------------------------------------------------------------------
# # Logging setup
# # ---------------------------------------------------------------------
# class Tee:
#     def __init__(self, *files): self.files = files
#     def write(self, obj):
#         for f in self.files:
#             f.write(obj); f.flush()
#     def flush(self):
#         for f in self.files: f.flush()

# timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
# SCRIPT_DIR = Path(__file__).resolve().parent.parent
# LOG_DIR = SCRIPT_DIR / "logs"
# LOG_DIR.mkdir(exist_ok=True)
# LOG_FILE = LOG_DIR / f"summary_log_{timestamp}.txt"
# log_f = open(LOG_FILE, "w", encoding="utf-8")
# sys.stdout = Tee(sys.__stdout__, log_f)
# sys.stderr = Tee(sys.__stderr__, log_f)

# load_dotenv()
# client = OpenAI()

# # ---------------------------------------------------------------------
# # Config
# # ---------------------------------------------------------------------
# TXT_DIR      = SCRIPT_DIR / "data" / "extracted_text" / "test4"
# IMAGES_DIR   = SCRIPT_DIR / "data" / "extracted_image" / "test4"
# PROMPT_PATH  = SCRIPT_DIR / "prompts" / "prompt_test.txt"
# MODEL        = "gpt-5"
# MAX_TOKENS   = 20000
# MAX_RETRIES  = 3

# OUTPUT_DIR   = SCRIPT_DIR / "data" / "summarize_to_doc_output"
# OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
# OUTPUT_DOCX  = OUTPUT_DIR / "Paper_Summary_Matched.docx"

# # ---------------------------------------------------------------------
# # Helpers
# # ---------------------------------------------------------------------
# def truncate_text(text: str, limit: int = 30_000) -> str:
#     return text if len(text) <= limit else text[:limit] + "\n\n[Text truncated for LLM]"

# def load_prompt() -> str:
#     return PROMPT_PATH.read_text(encoding="utf-8")

# # ---------------------------------------------------------------------
# # LLM Call (GPT-5 robust version)
# # ---------------------------------------------------------------------
# def call_llm(prompt: str) -> str:
#     """
#     Call GPT-5 and guarantee a valid JSON output string like:
#         {"SummaryDoc": "<document text>"}
#     Handles empty responses, retries, and GPT-5's strict formatting rules.
#     """

#     system_msg = {
#         "role": "system",
#         "content": (
#             "You are an academic summarizer and document restorer. "
#             "You MUST output only valid JSON with this schema:\n"
#             "{ \"SummaryDoc\": \"<full reconstructed academic document text including figure markers>\" }"
#         ),
#     }

#     user_msg = {"role": "user", "content": prompt}

#     for attempt in range(MAX_RETRIES):
#         try:
#             resp = client.chat.completions.create(
#                 model=MODEL,
#                 messages=[system_msg, user_msg],
#                 max_completion_tokens=MAX_TOKENS,
#                 response_format={"type": "json_object"},
#             )

#             print("LLM response repr (short):", repr(resp)[:2000])
#             choice = resp.choices[0]

#             content = getattr(choice.message, "content", None)
#             if not content:
#                 print("⚠️ Empty content field from LLM.")
#                 continue

#             # Clean and return JSON text
#             cleaned = clean_raw(content)
#             return cleaned.strip()

#         except Exception as e:
#             print(f"Attempt {attempt+1}/{MAX_RETRIES} failed: {e}")
#             if attempt == MAX_RETRIES - 1:
#                 raise
#             time.sleep(2 ** attempt)

#     return ""

# # ---------------------------------------------------------------------
# # Clean output text before JSON parsing
# # ---------------------------------------------------------------------
# def clean_raw(raw: str) -> str:
#     raw = raw.strip()
#     if raw.startswith("```"):
#         raw = raw.split("```", 1)[1].rsplit("```", 1)[0]
#     if raw.lower().startswith("json"):
#         raw = raw[4:].lstrip()
#     return raw.strip()

# # ---------------------------------------------------------------------
# # Insert image + caption
# # ---------------------------------------------------------------------
# def insert_image(doc: Document, img_path: Path, caption: str):
#     try:
#         p_img = doc.add_paragraph()
#         p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         run = p_img.add_run()
#         run.add_picture(str(img_path), width=Inches(5.5))

#         p_cap = doc.add_paragraph(caption, style="Caption")
#         p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         for run_cap in p_cap.runs:
#             run_cap.italic = True
#         doc.add_paragraph()
#     except Exception as e:
#         print(f"Could not insert image {img_path.name}: {e}")

# def ensure_caption_style(doc: Document):
#     if "Caption" not in doc.styles:
#         style = doc.styles.add_style("Caption", 1)
#         style.font.name = "Times New Roman"
#         style.font.size = Pt(10)
#         style.font.italic = True

# # ---------------------------------------------------------------------
# # Build DOCX
# # ---------------------------------------------------------------------
# def create_docx(summary_text: str, output_path: Path, images_dir: Path):
#     doc = Document()
#     normal = doc.styles["Normal"]
#     normal.font.name = "Times New Roman"
#     normal.font.size = Pt(12)

#     ensure_caption_style(doc)

#     image_map = {
#         p.name.lower(): p for p in images_dir.iterdir()
#         if p.suffix.lower() in {".png", ".jpg", ".jpeg", ".tif", ".bmp", ".svg"}
#     }

#     print(f"Found {len(image_map)} images in {images_dir}")

#     blocks = [b.strip() for b in summary_text.split("\n\n") if b.strip()]

#     for block in blocks:
#         line = block.strip()

#         if line.startswith("###"):
#             doc.add_heading(line.lstrip("# ").strip(), level=3)
#         elif line.startswith("##"):
#             doc.add_heading(line.lstrip("# ").strip(), level=2)
#         elif line.startswith("#"):
#             doc.add_heading(line.lstrip("# ").strip(), level=1)
#         elif line.startswith("[[FIGURE:"):
#             try:
#                 inner = line.strip("[]").replace("FIGURE:", "").strip()
#                 parts = [p.strip() for p in inner.split("|")]
#                 filename = parts[0]
#                 caption = " | ".join(parts[1:])
#                 img_key = Path(filename).name.lower()

#                 if img_key in image_map:
#                     insert_image(doc, image_map[img_key], caption)
#                     print(f"Inserted image: {filename}")
#                 else:
#                     print(f"Image NOT FOUND: {filename}")
#                     p = doc.add_paragraph(f"[Image missing: {filename}] {caption}", style="Normal")
#                     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             except Exception as e:
#                 print(f"Figure parsing error: {e}\n   Block: {line}")
#                 doc.add_paragraph(line, style="Normal")
#         else:
#             doc.add_paragraph(line, style="Normal")

#     doc.save(output_path)
#     print(f"DOCX saved → {output_path}")

# # ---------------------------------------------------------------------
# # Main
# # ---------------------------------------------------------------------
# def main():
#     print("\n=== Summarization with Images ===")
#     base_prompt = load_prompt()

#     txt_files = sorted(TXT_DIR.glob("*.txt"))
#     if not txt_files:
#         print(f"No .txt files in {TXT_DIR}")
#         return

#     combined = "\n\n".join(
#         f"--- FILE: {f.name} ---\n{truncate_text(f.read_text(encoding='utf-8'))}"
#         for f in txt_files
#     )

#     prompt = base_prompt.replace("<<<DOCUMENT_TEXT>>>", combined)
#     print(f"Calling LLM ({MODEL}) …")

#     raw = call_llm(prompt)

#     if not raw:
#         print("⚠️ Empty response from LLM. Aborting.")
#         return

#     try:
#         parsed = json.loads(raw)
#     except json.JSONDecodeError:
#         cleaned = clean_raw(raw)
#         try:
#             parsed = json.loads(cleaned)
#         except json.JSONDecodeError:
#             print("❌ Final JSON decode failed. Dumping first 4000 chars:\n")
#             print(raw[:4000])
#             return

#     summary = parsed.get("SummaryDoc", "").strip()
#     if not summary:
#         print("⚠️ LLM returned empty SummaryDoc.")
#         return

#     create_docx(summary, OUTPUT_DOCX, IMAGES_DIR)
#     print("\n✅ All done!")
#     print(f"Log: {LOG_FILE}")

# if __name__ == "__main__":
#     main()
