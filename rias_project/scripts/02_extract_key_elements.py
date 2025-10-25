import json
import os
import time
import base64
from pathlib import Path
from typing import Union, Iterable, Optional
from tqdm import tqdm
from docx import Document
from docx.shared import Inches

from openai import OpenAI
from dotenv import load_dotenv


# ------------------------------------------------------------------
# Load .env (API key, etc.)
# ------------------------------------------------------------------
load_dotenv()  # Looks for .env in cwd or parent dirs


def encode_image(image_path: Path) -> str:
    """Encode image to base64 string for OpenAI API."""
    with open(image_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def summarize_document_with_images(
    txt_path: Union[str, Path],
    output_dir: Union[str, Path],
    prompt_file: Union[str, Path],
    *,
    model: str = "gpt-4o-mini",
    max_tokens: int = 1200,
    temperature: float = 0.0
) -> Path:
    """
    Summarize a document using:
      - .txt file (full text)
      - Images in the **same folder** as the .txt file

    Output: .docx with summary + embedded images

    Args:
        txt_path: Path to the extracted .txt file
        output_dir: Base folder for output (per-doc subfolder created)
        prompt_file: Path to your custom prompt .txt
        model, max_tokens, temperature: OpenAI settings

    Returns:
        Path to generated .docx file
    """
    txt_path = Path(txt_path)
    output_dir = Path(output_dir)
    prompt_file = Path(prompt_file)

    # ------------------- Validation -------------------
    if not txt_path.exists():
        raise FileNotFoundError(f"Text file not found: {txt_path}")
    if not prompt_file.exists():
        raise FileNotFoundError(f"Prompt file not found: {prompt_file}")

    # Use the **same folder** as the .txt file for images
    images_dir = txt_path.parent
    if not any(images_dir.glob("*.png")) and not any(images_dir.glob("*.jpg")) and not any(images_dir.glob("*.jpeg")):
        print(f"Warning: No images found in {images_dir}")

    doc_stem = txt_path.stem
    doc_output_dir = output_dir / doc_stem
    doc_output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = doc_output_dir / f"{doc_stem}_summary.docx"

    # ------------------- Load Data -------------------
    full_text = txt_path.read_text(encoding="utf-8")
    prompt_template = prompt_file.read_text(encoding="utf-8")

    image_files = (
        sorted(images_dir.glob("*.png")) +
        sorted(images_dir.glob("*.jpg")) +
        sorted(images_dir.glob("*.jpeg"))
    )

    # ------------------- Prepare Messages -------------------
    messages = [{"role": "system", "content": "You are a precise document analyst. Use text and image layout to summarize accurately."}]

    # Truncate text safely
    text_content = full_text[:30000]
    if len(full_text) > 30000:
        text_content += "\n\n[Text truncated for length]"

    # Add images with positional context
    image_contexts = []
    for idx, img_path in enumerate(image_files, 1):
        b64 = encode_image(img_path)
        section = "top" if idx <= len(image_files)//3 else "middle" if idx <= 2*len(image_files)//3 else "bottom"
        caption = f"Image {idx} ({img_path.name}) - appears in the {section} section."

        image_contexts.extend([
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
            {"type": "text", "text": caption}
        ])

    user_content = [
        {"type": "text", "text": prompt_template.replace("<<<DOCUMENT_TEXT>>>", text_content)},
        *image_contexts
    ]
    messages.append({"role": "user", "content": user_content})

    # ------------------- Call OpenAI -------------------
    print(f"Calling {model} for {doc_stem}...")
    client = OpenAI()  # Uses OPENAI_API_KEY from .env

    try:
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            max_tokens=max_tokens,
            temperature=temperature
        )
        summary_text = response.choices[0].message.content.strip()
    except Exception as e:
        summary_text = f"[ERROR during LLM call: {str(e)}]"

    # ------------------- Generate .docx -------------------
    doc = Document()
    doc.add_heading(f"Summary: {doc_stem}", 0)

    for paragraph in summary_text.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    if image_files:
        doc.add_page_break()
        doc.add_heading("Reference Images", level=1)
        for idx, img_path in enumerate(image_files, 1):
            try:
                doc.add_paragraph(f"Image {idx}: {img_path.name}", style="Caption")
                doc.add_picture(str(img_path), width=Inches(5.5))
                doc.add_paragraph()
            except Exception as e:
                doc.add_paragraph(f"[Failed to embed {img_path.name}: {e}]", style="Caption")

    doc.save(docx_path)
    print(f"Saved: {docx_path}")
    return docx_path


# ===================================================================
# BATCH PROCESSOR (now uses txt_path.parent as images_dir)
# ===================================================================
def process_documents_batch(
    txt_dir: Union[str, Path],
    output_dir: Union[str, Path],
    prompt_file: Union[str, Path],
    *,
    pdf_names: Optional[Union[str, Iterable[str]]] = None,
    pattern: Optional[str] = None,
    delay: float = 1.0
) -> None:
    """
    Process all .txt files in txt_dir.
    Images are automatically loaded from the same folder as each .txt.
    """
    txt_dir = Path(txt_dir)
    output_dir = Path(output_dir)

    # Find .txt files where folder name == file stem
    txt_files = {
        p.parent.name: p
        for p in txt_dir.rglob("*.txt")
        if p.parent.name == p.stem
    }

    if pdf_names:
        if isinstance(pdf_names, str):
            pdf_names = [pdf_names]
        txt_files = {k: v for k, v in txt_files.items() if k in pdf_names}

    if pattern:
        import fnmatch
        txt_files = {k: v for k, v in txt_files.items() if fnmatch.fnmatch(k, pattern)}

    if not txt_files:
        print("No documents matched.")
        return

    print(f"Processing {len(txt_files)} document(s)...\n")

    for name, txt_path in tqdm(txt_files.items(), desc="Summarizing", unit="doc"):
        summarize_document_with_images(
            txt_path=txt_path,
            output_dir=output_dir,
            prompt_file=prompt_file
        )
        time.sleep(delay)


# ===================================================================
# EXAMPLE USAGE
# ===================================================================
if __name__ == "__main__":
    load_dotenv()

    SCRIPT_DIR = Path(__file__).resolve().parent.parent

    # 1. Single document
    summarize_document_with_images(
        txt_path=SCRIPT_DIR / "data/extracted_text/Invoice_001/Invoice_001.txt",
        output_dir=SCRIPT_DIR / "data/summarized_docs",
        prompt_file=SCRIPT_DIR / "prompts/summary_with_images.txt"
    )

    # 2. Batch process
    # process_documents_batch(
    #     txt_dir=SCRIPT_DIR / "data/extracted_text",
    #     output_dir=SCRIPT_DIR / "data/summarized_docs",
    #     prompt_file=SCRIPT_DIR / "prompts/summary_with_images.txt",
    #     pattern="*invoice*"
    # )